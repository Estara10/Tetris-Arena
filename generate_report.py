#!/usr/bin/env python3
"""生成优化后的专利说明书 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=None):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_cover_page(doc):
    """添加封面页 - 与原稿保持一致"""
    # 空行
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(60)

    # Harbin Institute of Technology
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Harbin Institute of Technology")
    run.font.size = Pt(9)
    run.font.bold = True
    p.paragraph_format.first_line_indent = Cm(0)

    # 课程设计说明书
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("课程设计说明书")
    run.font.size = Pt(48)
    run.font.bold = True
    p.paragraph_format.first_line_indent = Cm(0)

    # 空行
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

    # 课程信息
    info_lines = [
        ("课程名称：", "智能控制项目实践"),
        ("", "课程项目报告"),
        ("设计题目：", "一种基于深度Q学习的同屏"),
        ("", "俄罗斯方块对抗系统及方法"),
        ("院    系：", "未来技术学院"),
        ("年    级：", "2023级本科生"),
        ("设 计 者：", "商喜庆"),
        ("学    号：", "2023112491"),
        ("指导教师：", "张淼"),
        ("设计时间：", "2026年春季学期"),
    ]

    for label, value in info_lines:
        p = doc.add_paragraph()
        if label:
            p.paragraph_format.first_line_indent = Cm(1.2)
        else:
            p.paragraph_format.first_line_indent = Cm(2.4)
        run = p.add_run(label + value)
        run.font.size = Pt(18)
        if not label:
            run.font.bold = False

    # 空行
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

    # 哈尔滨工业大学
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("哈尔滨工业大学")
    run.font.size = Pt(16)
    p.paragraph_format.first_line_indent = Cm(0)

    # 分页
    doc.add_page_break()


def add_task_sheet(doc):
    """添加任务书 - 与原稿保持一致"""
    # 空行
    p = doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("智能控制项目实践——任务书")
    run.font.size = Pt(18)
    run.font.bold = True
    set_paragraph_spacing(p, before=0, after=12)

    # 1. 课程报告要求
    p = doc.add_paragraph()
    run = p.add_run("1．课程报告要求：")
    run.font.size = Pt(12)
    set_paragraph_spacing(p, before=0, after=4)

    p = doc.add_paragraph()
    run = p.add_run("结合实践项目及课程所讲述的分析处理、特征提取、深度网络、强化学习等方面内容，以自身在项目研发工作中的软件课题/方法实现为课程报告内容，按照专利撰写规范，题目自拟，撰写一份发明专利申请材料，重点涉及专利文件最核心的说明书（含附图）的撰写。")
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0.74)
    set_paragraph_spacing(p, before=0, after=6)

    # 2. 任务工作量
    p = doc.add_paragraph()
    run = p.add_run("2．任务工作量：")
    run.font.size = Pt(12)
    set_paragraph_spacing(p, before=0, after=4)

    p = doc.add_paragraph()
    run = p.add_run("说明书撰写应包含以下工作：")
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0.74)
    set_paragraph_spacing(p, before=0, after=4)

    workload_items = [
        "（1）技术背景介绍；",
        "（2）主要涉及方案；",
        "（3）发明内容的有益效果分析；",
        "（4）给出必要的说明书附图，以增强说明书的可读性；",
        "（5）设计方案的实施例。",
    ]
    for item in workload_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        set_paragraph_spacing(p, before=0, after=2)

    # 3. 评分标准
    p = doc.add_paragraph()
    run = p.add_run("3．评分标准——从以下四方面综合考虑（占比相同）")
    run.font.size = Pt(12)
    set_paragraph_spacing(p, before=8, after=4)

    scoring_items = [
        "（1）格式规范性（25分）",
        "（2）内容原创性（25分）",
        "（3）说明书内容是否明确指明了实际解决的技术问题，即为获得更好的技术效果而需要对最接近现有技术进行改进的技术任务，是否明确指明了区别特征，并且是应用该特征来解决实际问题或困难的（25分）",
        "（4）程序代码的完善度，请将源程序（含打包版）和报告一并提交（25分）",
    ]
    for item in scoring_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        set_paragraph_spacing(p, before=0, after=2)

    # 4. 工作计划
    p = doc.add_paragraph()
    run = p.add_run("4．工作计划")
    run.font.size = Pt(12)
    set_paragraph_spacing(p, before=8, after=4)

    plan_items = [
        "（1）理论阶段：掌握课程所讲内容，树立创新观念，结合老师给定大方向自拟开发项目题目并明确说明书背景。",
        "（2）实践阶段：实现项目基本功能、结合深度学习或其他智能算法，自拟仿真题目，通过计算机仿真，完成发明内容的主要步骤。",
        "（3）课程最后1节（项目验收）：按照模板/样例，撰写开发文档，打包源代码，课上展示Demo，介绍开发历程，开发文档连同程序代码一同拷给老师。",
        "（4）课程结束后1周（模拟专利）：撰写发明实施例，完善有益效果分析、说明书附图等环节，同时写明如何应用发明的区别特征来解决实际问题或困难的。",
    ]
    for item in plan_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        set_paragraph_spacing(p, before=0, after=2)

    # 教师签字
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("教师签字：___________")
    run.font.size = Pt(12)
    set_paragraph_spacing(p, before=12, after=6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("（本页无需填写，所提交课程报告保留本页）")
    run.font.size = Pt(12)
    run.font.bold = True

    # 分页
    doc.add_page_break()


def add_section_header(doc, text):
    """添加章节标题（如 （一）技术领域）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    set_paragraph_spacing(p, before=12, after=6)
    return p


def add_body_paragraph(doc, text, first_indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    set_paragraph_spacing(p, before=0, after=4)
    return p


def add_step_header(doc, text):
    """添加步骤标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.first_line_indent = Cm(0.74)
    set_paragraph_spacing(p, before=8, after=4)
    return p


def add_formula(doc, text):
    """添加公式（居中）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.italic = True
    set_paragraph_spacing(p, before=4, after=4)
    return p


def add_patent_body(doc):
    """添加专利说明书正文（优化版）"""

    # ===== 标题 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("一种基于深度Q学习的同屏俄罗斯方块对抗系统及方法")
    run.font.size = Pt(15)
    run.font.bold = True
    set_paragraph_spacing(p, before=0, after=18)

    # ===== （一）技术领域 =====
    add_section_header(doc, "（一）技术领域")

    add_body_paragraph(doc,
        "本发明涉及计算机游戏智能控制、人机交互与强化学习决策技术领域，"
        "具体涉及一种基于深度Q学习的同屏俄罗斯方块对抗系统及方法。"
        "该系统面向俄罗斯方块实时竞技场景，通过棋盘状态特征提取、候选落点评估、"
        "深度Q网络价值预测以及自适应旁路降级机制，实现人类玩家与智能体之间的实时对抗控制。"
    )

    # ===== （二）背景技术 =====
    add_section_header(doc, "（二）背景技术")

    add_body_paragraph(doc,
        "俄罗斯方块是一类典型的离散状态、连续决策、实时反馈型益智游戏。"
        "由于其规则明确、状态空间动态变化明显、动作选择对后续局面影响较大，"
        "俄罗斯方块常被用于人工智能算法验证、智能控制实验和人机博弈系统设计。"
        "在常规游戏实现中，系统通常只包含方块生成、旋转、平移、下落、消行和计分等基础功能，"
        "玩法多以单人模式为主，缺少高实时性的人机对抗机制，也难以体现智能控制算法在动态交互场景中的应用价值。"
    )

    add_body_paragraph(doc,
        "在俄罗斯方块智能决策算法方面，现有方法主要包括固定启发式算法、搜索算法以及基于神经网络的学习算法。"
        "固定启发式算法通常根据孔洞数量、堆叠高度、表面平整度等人工指标对落点进行评分，"
        "具有实现简单、运行速度快的优点，但策略较为固定，难以适应复杂对抗场景。"
        "搜索算法能够在一定范围内遍历候选动作并选择局部较优方案，"
        "但随着棋盘状态和候选动作数量增加，计算开销会显著上升，不利于在实时游戏主循环中稳定运行。"
        "基于深度学习的算法能够通过训练获得较强的状态评估能力，"
        "但若直接应用于实时对抗系统，仍可能受到模型推理延迟、状态适配性和主线程阻塞等问题影响。"
    )

    add_body_paragraph(doc,
        "在同屏或双模式竞技场景中，系统不仅需要完成俄罗斯方块基本规则，"
        "还需要同步处理玩家输入、AI决策、棋盘渲染、碰撞检测和对局结算等任务。"
        "若AI在每一帧中执行过深的前瞻搜索，可能导致游戏画面卡顿，影响玩家操作体验；"
        "若AI仅采用简单规则，又难以体现智能控制和强化学习方法的优势。"
        "因此，有必要设计一种兼顾实时性、可运行性和智能决策能力的俄罗斯方块对抗系统，"
        "使其既能体现深度Q学习在候选状态评估中的作用，又能通过工程化调度机制保障游戏运行流畅。"
    )

    add_body_paragraph(doc,
        "基于上述问题，本文提出一种基于深度Q学习的同屏俄罗斯方块对抗系统及方法。"
        "该方法通过低维棋盘特征表示、候选落点状态评估、深度Q网络价值预测和自适应旁路降级机制，"
        "将智能决策过程嵌入Pygame游戏主循环中，"
        "从而在课程项目实现中验证强化学习算法与实时交互系统结合的可行性。"
    )

    # ===== （三）发明内容 =====
    add_section_header(doc, "（三）发明内容")

    add_body_paragraph(doc,
        "本发明的目的在于提供一种基于深度Q学习的俄罗斯方块智能对抗方法，"
        "用于解决传统启发式策略适应性不足、深度前瞻搜索计算开销较大以及实时游戏对抗中容易出现主线程阻塞的问题。"
        "该方法在保留俄罗斯方块基础玩法的前提下，引入候选落点状态评估和自适应旁路降级机制，"
        "使AI能够根据不同对局模式选择合适的决策强度，从而兼顾智能性与实时性。"
    )

    add_body_paragraph(doc, "为实现上述目的，本发明采用如下技术方案。")

    # ---- 步骤一 ----
    add_step_header(doc, "步骤一：构建轻量级棋盘状态特征空间")

    add_body_paragraph(doc,
        "系统实时获取当前俄罗斯方块棋盘的二维网格数据，并结合当前下落方块及候选落点信息，"
        "对执行候选动作后的棋盘状态进行模拟。针对每一个候选状态，系统提取行消除数量、棋盘空洞数量、"
        "列高度总和以及表面起伏程度等关键特征，并将其组织为低维状态特征向量，用于后续价值评估。"
    )

    add_body_paragraph(doc,
        "状态特征向量可表示为：St = [fL, fH, fA, fB]^T，"
        "其中 fL 表示当前候选动作可消除的行数，fH 表示棋盘中的空洞数量，"
        "fA 表示各列高度之和，fB 表示棋盘表面粗糙度。"
        "表面粗糙度由相邻列高度差的绝对值之和计算："
    )
    add_formula(doc, "fB = Σ(i=1 to W-1) |h_i - h_(i+1)|")
    add_body_paragraph(doc,
        "其中 W 为棋盘总列数，h_i 为第 i 列的堆叠高度。"
        "通过上述特征提取方式，系统避免直接处理高维图像输入，将复杂棋盘状态转化为适合实时计算的数值特征。"
    )

    # ---- 步骤二 ----
    add_step_header(doc, "步骤二：构建动作价值反馈评估网络")

    add_body_paragraph(doc,
        "系统将候选动作执行后的状态特征向量输入预训练的深度Q网络（Deep Q-Network, DQN）。"
        "该网络由多层全连接结构（隐藏层维度为 128、128、64）和非线性激活单元组成，"
        "用于估计候选落点状态的价值。对于当前状态下的每一个候选动作 a，"
        "系统先模拟该动作落子后的后继状态 S_(t+1)^(a)，再由网络输出对应的价值评估结果 V(S_(t+1)^(a); θ)，"
        "其中 θ 表示深度Q网络的参数。"
    )

    add_body_paragraph(doc,
        "系统根据该价值结果对候选动作进行排序，并优先选择预期收益较高、局面风险较低的落点方案。"
        "与单纯依赖人工规则评分的方法相比，该方式能够利用训练过程中获得的状态价值经验，"
        "提高AI决策的适应性。在模型推理模式下，系统先由启发式方法生成候选动作集（通常包含8个候选方案），"
        "再由深度Q网络对候选集进行价值重排，从而兼顾计算效率与决策质量。"
    )

    # ---- 步骤三 ----
    add_step_header(doc, "步骤三：构建基于前瞻权重因子的自适应旁路降级机制")

    add_body_paragraph(doc,
        "为降低实时对抗场景中的计算压力，系统设置前瞻权重因子 lookahead_weight，记为 α。"
        "当系统允许进行较深层次预测时，目标评估函数可表示为："
    )
    add_formula(doc, "E(a) = V(S_(t+1)^(a); θ) + α · max(a'∈A') V(S_(t+2)^(a,a'); θ)")
    add_body_paragraph(doc,
        "式中，A' 表示下一步候选动作集合，S_(t+2)^(a,a') 表示连续执行当前候选动作和下一候选动作后的预测状态。"
        "该方式能够在一定程度上考虑后续方块对整体局面的影响。"
    )

    add_body_paragraph(doc,
        "当系统检测到当前模式对响应速度要求较高，或前瞻权重 α ≤ 0 时，"
        "控制器主动触发旁路降级逻辑，跳过下一层候选状态遍历，"
        "仅保留当前候选状态的价值评估："
    )
    add_formula(doc, "E_bypass(a) = V(S_(t+1)^(a); θ)")
    add_body_paragraph(doc,
        "该机制将原本可能涉及二级候选展开的计算过程简化为单层候选评估，"
        "将单步动作寻优的时间复杂度从 O(|A| × |A'|) 降至 O(|A|)，"
        "计算耗时由百毫秒级压缩至毫秒级常数范围，从而在算法层面有效缓解高频竞技要求下的主线程阻塞与渲染卡顿问题。"
        "对于高频对抗模式（如经典双屏对战），该机制能够优先保障响应速度；"
        "对于需要更充分规划的模式（如同屏竞技），则可保留一定前瞻能力。"
    )

    # ---- 步骤四 ----
    add_step_header(doc, "步骤四：执行基于多模式隔离的竞技帧同步调度")

    add_body_paragraph(doc,
        "系统根据当前对局模式配置AI控制器参数，并在Pygame主循环中统一调度玩家输入、"
        "AI决策、棋盘更新和画面渲染。玩家输入通过非阻塞事件机制读取，"
        "AI决策根据设定的动作间隔（action_interval_ms）和前瞻权重（α）生成控制指令，"
        "双方状态在固定帧率（60 FPS）下同步更新。"
    )

    add_body_paragraph(doc,
        "在课程项目实现中，系统支持经典双屏对战和同屏竞技两种模式。"
        "经典模式下，玩家与AI在左右两个独立棋盘中使用同步方块序列进行对战，"
        "前瞻权重 α 设为0以触发旁路降级，保障输入响应流畅；"
        "同屏竞技模式下，玩家与AI在同一25×25共享棋盘空间内进行推挤式竞争，"
        "系统可配置较高的前瞻权重以保留模型的前瞻评估能力。"
        "不同模式通过调整AI动作间隔、错误率、模型启用策略和前瞻权重，"
        "实现规则隔离和策略适配，从而在保持交互流畅的同时展示深度Q学习方法在游戏智能控制中的应用效果。"
    )

    # ---- 有益效果 ----
    add_step_header(doc, "有益效果")

    add_body_paragraph(doc, "与现有技术相比，本发明具有以下有益效果。")

    add_body_paragraph(doc,
        "第一，本发明能够提升AI决策过程的实时性。"
        "通过设置基于前瞻权重的自适应旁路降级机制，系统可在高频交互模式下跳过计算量较大的二级候选遍历，"
        "将AI决策控制在更短时间内完成，从而降低主线程阻塞和画面卡顿的概率。"
        "实验表明，在经典模式下采用旁路降级后，AI单次寻优耗时由约141.684毫秒降至约2.075毫秒，"
        "降幅达98.5%以上，显著缓解了实时对抗中的延迟问题。"
    )

    add_body_paragraph(doc,
        "第二，本发明增强了同一AI决策框架在不同模式下的适配能力。"
        "系统不需要为每一种对局模式重新设计完整算法，"
        "而是通过调整前瞻权重、动作间隔和控制参数，"
        "在“快速响应”和“适度前瞻”之间进行切换，"
        "使AI策略更适合课程项目中的多模式俄罗斯方块对抗场景。"
    )

    add_body_paragraph(doc,
        "第三，本发明将深度Q学习模型与工程化游戏系统进行了结合。"
        "候选落点状态先经过特征提取和模拟，再输入深度Q网络进行价值评估，"
        "既避免了直接处理高维图像输入带来的计算负担，"
        "也能够体现强化学习模型在状态评估中的作用。"
        "该混合决策架构使AI同时具备启发式方法的稳定性和学习模型的适应性。"
    )

    add_body_paragraph(doc,
        "第四，本发明具有较好的课程实践可验证性。"
        "系统包含可运行的Pygame游戏界面、AI控制模块、训练曲线、模型推理和实机对抗效果图，"
        "能够从代码实现、实验结果和演示界面三个方面说明方案的可行性。"
        "同时，系统提供的三种AI难度等级（简单/普通/困难）和可视化对抗界面，"
        "便于在教学和演示场景中直观展示算法效果。"
    )

    # ===== （四）附图说明 =====
    add_section_header(doc, "（四）附图说明")

    add_body_paragraph(doc,
        "为了更清楚地说明本发明实施例中的技术方案，下面结合附图对本发明的主要结构和流程进行说明。"
        "所列附图用于辅助理解本课程项目中的系统设计与实现方式，并不构成对具体实现形式的限制。"
    )

    add_body_paragraph(doc,
        "图1为本发明实施例提供的俄罗斯方块智能决策方法整体流程示意图，"
        "展示了状态获取、候选动作生成、价值评估、旁路判断和指令输出的基本过程。"
    )
    add_body_paragraph(doc,
        "图2为本发明实施例中特征提取机制的网格状态转化示意图，"
        "展示了如何从棋盘网格中提取高度、空洞、消行和表面起伏等低维特征。"
    )
    add_body_paragraph(doc,
        "图3为本发明实施例中自适应降级旁路调度机制的逻辑判断流程图，"
        "展示了系统根据前瞻权重决定是否执行下一层候选状态预测的过程。"
    )
    add_body_paragraph(doc,
        "图4为本发明实施例中深度Q网络模型训练过程的收敛曲线示意图，"
        "用于说明模型训练过程中损失值（Loss）和消行数（Lines Cleared）等评估指标的变化情况。"
    )
    add_body_paragraph(doc,
        "图5为本发明实施例系统在多模式竞技场景下的实际运行界面效果图，"
        "包括：(a) 经典双屏对战模式下的人机博弈界面；"
        "(b) 同屏竞技模式下的共享棋盘对抗界面。"
    )

    # 图片建议
    add_body_paragraph(doc,
        "【图片完善建议】"
        "（1）图1流程图建议使用标准流程图符号（圆角矩形/菱形/箭头），标注每个模块对应的代码文件；"
        "（2）图2特征提取示意图可增加数值标注示例，展示从二维网格到一维特征向量的具体转化过程；"
        "（3）图3旁路判断流程图建议使用决策树形式呈现，明确标注 α > 0 和 α ≤ 0 两条分支；"
        "（4）图4训练曲线建议同时绘制 Loss 曲线和 Lines Cleared 曲线，并标注课程学习阶段切换点（1200 episodes）；"
        "（5）图5界面截图建议截取亮色主题版本（当前已实现），并添加文字标注说明各UI区域功能。",
        first_indent=False
    )

    # ===== （五）具体实施方式 =====
    add_section_header(doc, "（五）具体实施方式")

    add_body_paragraph(doc,
        "下面结合附图和课程项目实现，对本发明实施例中的技术方案进行清楚、完整的描述。"
        "应当理解，以下实施例用于说明本系统的实现方式，并不表示本发明只能采用该一种具体实现。"
    )

    # ---- 技术栈 ----
    add_step_header(doc, "技术栈")

    add_body_paragraph(doc,
        "本实施例所涉及的主要技术组件如下：",
        first_indent=True,
    )
    add_body_paragraph(doc,
        "游戏引擎：Pygame（2D 渲染、事件处理、60 FPS 主循环调度）",
        first_indent=False,
    )
    add_body_paragraph(doc,
        "AI 框架：PyTorch + Deep Q-Network（三层全连接 128/128/64，23 维特征输入）",
        first_indent=False,
    )
    add_body_paragraph(doc,
        "训练策略：Prioritized Experience Replay + 课程学习 + Teacher Forcing",
        first_indent=False,
    )
    add_body_paragraph(doc,
        "语言与数据：Python 3.10 + NumPy",
        first_indent=False,
    )
    add_body_paragraph(doc,
        "通信机制：Pygame 事件系统（键盘输入非阻塞轮询，AI 决策帧同步调度）",
        first_indent=False,
    )
    add_body_paragraph(doc,
        "存储：PyTorch .pt 模型文件",
        first_indent=False,
    )
    add_body_paragraph(doc,
        "工具链：VS Code + Git + python-docx",
        first_indent=False,
    )

    # 结合图1
    add_body_paragraph(doc,
        "结合图1（智能决策方法整体流程示意图），本发明提供了一种基于深度Q网络与自适应降级机制的"
        "俄罗斯方块智能处理架构。系统在初始化阶段配置Pygame渲染主循环，"
        "并加载预训练的DQN神经网络权重（.pt模型文件，存储于models/next_state_v3/目录下）。"
        "在游戏对局的每一帧中，首先由环境感知模块获取当前二维网格矩阵数据与下落方块（Tetromino）的绝对坐标集，"
        "然后遍历当前所有合法的平移与旋转组合空间，构成候选动作集，作为后续决策引擎的输入基础。"
        "系统还通过7-bag发牌机制确保玩家与AI使用相同的方块序列，保证对局的公平性。"
    )

    # 结合图2
    add_body_paragraph(doc,
        "结合图2（特征提取网格状态转化示意图），在动作评估前，本系统通过特征工程执行轻量级的数据降维。"
        "具体实施中，算法并非将原始棋盘网格直接输入卷积神经网络，而是构建了一套人工特征过滤管道。"
        "系统模拟方块在某个特定位置落底后的局面，"
        "并即时计算该虚拟局面的多项一维关键指标：包括各列天际线高度（Column Heights）、"
        "形成的孔洞数量（Holes）、相邻列的高度差绝对值之和（Bumpiness），以及模拟消除的行数（Cleared Lines）。"
        "上述特征通过next_state_features.py模块中的特征提取函数统一计算，"
        "将复杂的高维游戏状态压缩为23维低维特征向量，实现低延迟的实时计算。"
    )

    # 结合图4
    add_body_paragraph(doc,
        "结合图4（模型训练收敛曲线图），为了突破传统启发式规则容易陷入局部最优的局限，"
        "本发明采用了深度Q网络（Deep Q-Network）进行动作价值预测。"
        "训练过程采用Prioritized Experience Replay（优先级经验回放，α=0.6, β: 0.4→1.0）、"
        "课程学习（前1200 episodes为单人预训练阶段，逐步过渡到对抗训练）和"
        "Teacher Forcing等策略。从图4中的Loss曲线平滑下降并趋于稳定，"
        "以及Lines Cleared曲线随训练进行产生跃升式增长可以看出，"
        "网络通过反向传播算法掌握了环境状态与未来预期收益（Q值）之间的非线性映射关系。"
        "在实机部署时，将上述降维后的特征向量输入该网络，网络即可输出各个候选落点状态的价值评分。"
    )

    # 结合图3
    add_body_paragraph(doc,
        "结合图3（自适应降级旁路调度逻辑流程图），本系统的核心创新点在于针对实时竞技卡顿（Stuttering）"
        "问题设计的自适应旁路降级机制。在实施过程中（参见ai_controller.py），"
        "对于包含未来状态深度预测（get_best_future_score方法）的模块增设了严格的条件旁路器。"
        "当系统检测到当前处于高频并发的经典双屏对战模式时，控制器层自动将前瞻权重因子 lookahead_weight 设为0；"
        "系统随后触发逻辑旁路，强制跳过对下一次落块环境的完全遍历预测过程，"
        "转而直接输出具有当前最大基础Q值（或启发式评分）的落点方案。"
        "此设计从根本上切断了由于状态树二次展开所导致的算力溢出，"
        "实验测得该降级机制将单次寻优耗时由141.684毫秒显著降至2.075毫秒，完全适应高频实时响应需求。"
    )

    # 结合图5
    add_body_paragraph(doc,
        "结合图5（多模式竞技场景下的实机博弈效果图），本方案最终落地为统一渲染循环下的双实体同步调度框架。"
        "在具体的游戏对战中，主界面的渲染、人类玩家键盘事件的非阻塞侦听与读取，"
        "以及AI的智能决策函数均在主流程中由固定刷新时钟（pygame.time.Clock.tick(60)）统一约束。"
        "同屏竞技模式下，系统还实现了活动方块间的碰撞推挤机制和垃圾行攻击/抵消系统，"
        "增强了对抗的策略深度和观赏性。AI的高速不卡顿推演保障了人类玩家在同一终端、"
        "同一UI视口下的平滑操作体验，成功构筑了流畅、智能且具有挑战性的人机博弈系统。"
    )

    # 视觉反馈系统
    add_step_header(doc, "视觉反馈增强")

    add_body_paragraph(doc,
        "为进一步提升人机交互的直观性和游戏体验的观赏性，本实施例还实现了一套轻量级视觉反馈系统，"
        "包含以下四项独立效果模块，均封装于effects.py视觉特效模块中，"
        "对游戏逻辑零侵入，仅依赖渲染层的事件触发接口。"
    )
    add_body_paragraph(doc,
        "（1）消行粒子爆发与屏幕震动：当任意实体完成行消除时，系统沿消除行所在像素位置"
        "向四周随机喷射30个彩色粒子（具备重力加速度与二次衰减透明度），"
        "同时根据消除行数触发分级屏幕震动（1行2像素振幅，4行10像素振幅，衰减周期280-420毫秒），"
        "强化消行操作的打击反馈感。"
    )
    add_body_paragraph(doc,
        "（2）飞字提示：消除行时从消除位置向上飘出浮动文字（如「TETRIS!」「COMBO x3」「+400」），"
        "含上升动画与透明度衰减，持续约1.4秒后自动消失，为玩家提供清晰的即时操作反馈。"
    )
    add_body_paragraph(doc,
        "（3）方块锁定闪光：每次方块落底锁定（Lock）时，"
        "在该方块所占据的所有单元格上叠加180毫秒白色高亮闪光并快速衰减，"
        "使玩家能够准确感知方块固化的瞬间位置。"
    )
    add_body_paragraph(doc,
        "（4）菜单淡入过渡：模式选择与AI等级选择界面在首次进入及切换时，"
        "执行500毫秒缓入淡出过渡动画（三次缓动函数），替代传统的瞬间画面跳变，"
        "提升UI交互的流畅度和完成度。"
    )
    add_body_paragraph(doc,
        "上述视觉增强模块合计约150行代码，均在渲染层独立运行，"
        "不影响AI决策逻辑与游戏核心状态机，体现了「表现层与逻辑层分离」的工程原则。"
    )

    # 实施例结论
    add_step_header(doc, "实施例结论")

    add_body_paragraph(doc,
        "依据本实施例的运行结果与训练曲线可知，所提出的基于深度Q网络与自适应旁路降级机制的"
        "俄罗斯方块智能决策方法能够稳定运行并完成实时人机对抗；"
        "模型训练过程中损失函数整体下降且评估指标趋于稳定，表明该方法具备有效的策略学习能力；"
        "在经典双屏对战模式下，采用前瞻权重旁路机制后，AI单帧决策耗时由约141.684毫秒降至约2.075毫秒，"
        "显著缓解主线程阻塞并消除对抗过程中的卡顿现象。"
        "综上，本方法在可运行性、对抗性与实时性方面均达到预期目标，"
        "在课程实践要求下具备较好的完整性、原创性和可展示性。"
    )

    # 分页后留空
    doc.add_page_break()


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 添加封面
    add_cover_page(doc)

    # 添加任务书
    add_task_sheet(doc)

    # 添加专利正文
    add_patent_body(doc)

    # 保存
    output_path = "/home/lidio/Tetris_Python/商喜庆_2026春季实践课程报告_优化版v2.docx"
    doc.save(output_path)
    print(f"优化版报告已保存至: {output_path}")


if __name__ == "__main__":
    main()
