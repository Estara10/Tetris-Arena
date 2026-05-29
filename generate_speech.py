#!/usr/bin/env python3
"""生成项目展示发言稿 docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(16)


def add_section(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(140, 140, 160)
    p.paragraph_format.space_after = Pt(4)


def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— — —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 180, 200)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


def main():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ==============================================
    add_title(doc, "俄罗斯方块 AI 对战 — 项目展示发言稿")
    add_note(doc, "预计时长：5–7 分钟 ｜ 配合游戏录屏播放 ｜ 重点：创新点+巧思")

    # ==============================================
    add_section(doc, "一、开场：我和大家做的是同一件事，但我用了不同的思路")

    add_body(doc,
        "咱们这门课大家做的都是俄罗斯方块 AI 对抗项目，所以方块怎么转、消行怎么算分"
        "这些基础的东西我就不讲了。我直接说我的项目跟别人不一样的地方。"
    )
    add_body(doc,
        "我觉得一个俄罗斯方块 AI 好不好，核心就两个问题：第一，AI 能不能做出聪明的决策；"
        "第二，AI 做决策的时候游戏画面会不会卡。"
        "多数方案要么只注重 AI 的智能性、把游戏跑得很卡，"
        "要么为了保证流畅用了太简单的规则、AI 没什么学习能力。"
        "我这个项目想同时解决这两个问题。"
    )

    add_divider(doc)

    # ==============================================
    add_section(doc, "二、我的核心创新：自适应旁路降级机制")
    add_note(doc, "【重点讲这段，这是最大的亮点】")

    add_body(doc,
        "先讲最大的创新点。我管它叫「自适应旁路降级」。"
    )
    add_body(doc,
        "AI 在决定方块放哪的时候，通常有两种做法：一种是只看当前这一步，"
        "好处是快，坏处是眼光短；另一种是往前多想一步——"
        "「如果我放这里，下一块也考虑进去会怎样」——好处是更聪明，"
        "但计算量会从一次决策变成几十次，主线程扛不住。"
    )
    add_body(doc,
        "我的巧思在于：不是一刀切地选快或选聪明，而是根据游戏模式自动切换。"
        "经典双屏对战模式节奏快、动作密集，我就让前瞻权重设为 0，直接跳过二级搜索，"
        "把 AI 寻优耗时控制在 2 毫秒以内——跟别的方案比，这个数字低了近 70 倍，画面完全不卡。"
        "而同屏竞技模式方块落得慢、更讲究布局，我就把前瞻权重拉高，"
        "让 AI 多算一步，充分评估后续方块的影响。"
    )
    add_body(doc,
        "这个设计的好处是：不需要为每种模式单独写一套 AI，"
        "一个框架、调一个参数，就能在「快」和「聪明」之间平滑切换。"
    )

    add_divider(doc)

    # ==============================================
    add_section(doc, "三、另一个巧思：启发式+深度Q网络混合决策")

    add_body(doc,
        "再讲 AI 本身的决策架构。大家可能有的用纯规则、有的用纯神经网络。"
        "我的做法是把两者揉在一起。"
    )
    add_body(doc,
        "先用启发式规则——就是根据洞的数量、表面平整度这些指标——快速筛出 8 个候选落点。"
        "这个步骤很快，几十行代码，但能覆盖绝大多数合理的选择。"
        "然后让预训练的深度 Q 网络在这 8 个候选里做精细排序，"
        "相当于启发式负责「海选」，神经网络负责「决赛」。"
    )
    add_body(doc,
        "这样做有两个好处：一是神经网络不用对所有可能位置逐一打分，计算量大幅降低；"
        "二是万一模型加载失败或者推理出错，系统自动回退到纯启发式模式，"
        "相当于自带了一个安全兜底。实际训练的时候，我还加了优先经验回放和课程学习，"
        "让模型多练那些它容易判断错的状态，前 1200 个 episode 先单人打基础再引入对抗，"
        "训练效率比均匀采样高不少。"
    )

    add_divider(doc)

    # ==============================================
    add_section(doc, "四、同屏竞技模式：不仅仅是多一个棋盘")

    add_body(doc,
        "除了经典的双屏对战，我还做了一个同屏竞技模式。"
        "这不是简单地把两个棋盘拼在一起——它有几个独特的设计。"
    )
    add_body(doc,
        "第一，碰撞推挤。玩家和 AI 的活动方块如果在水平方向撞上，会互相推动。"
        "系统会先计算一整条被推动的链条——A 推 B、B 推 C——如果链条末端不会撞墙，"
        "整条链就一起平移；如果末端撞墙，整个推动就无效。"
        "这给对抗增加了一层物理博弈的维度，你可以用方块把对手挤到不利的位置。"
    )
    add_body(doc,
        "第二，推送权机制。在 1v1 模式下，推送权不是谁都有的——"
        "当前拥有推送权的一方才能主动推挤对方，推完之后权力转移给被推的人。"
        "这就防止了某一方无限推挤，让博弈有来有回。"
    )
    add_body(doc,
        "第三，5 分钟倒计时制，时间到了按得分定胜负，不是传统的「谁先顶到头谁输」。"
        "这个设计让比赛更有紧张感，最后三十秒疯狂抢分的场面特别好看。"
    )
    add_body(doc,
        "另外 UI 上做了一些趣味细节：消行有黄色闪光特效，"
        "结束画面输了会弹出「人类一败涂地，菜就多练」，赢了显示"
        "「你的代码没有我的手速快」——纯粹为了好玩。"
    )

    add_divider(doc)

    # ==============================================
    add_section(doc, "五、录屏展示（配合解说）")
    add_note(doc, "【切换到录屏播放】")

    add_body(doc,
        "好，直接看录屏。这是经典双屏对战——左边我操作，右边是 AI。"
        "注意看 AI 落子的节奏，没有延迟感，这就是旁路降级在起作用。"
        "切换到困难难度之后，AI 消行效率明显提高，基本零失误。"
    )
    add_body(doc,
        "另外注意消行时的视觉反馈——粒子爆发特效从消除行向四周喷射，"
        "同时屏幕会有一个轻微的震动，消除行数越多震动越强，四行连消时最明显。"
        "消行位置会飘出得分文字和连击提示，像「TETRIS!」「COMBO x3」这样的飞字特效。"
        "这些细节让游戏的打击感和观赏性强了很多。"
    )
    add_note(doc, "【等录屏切换到同屏竞技部分】")

    add_body(doc,
        "这段是同屏竞技。注意左上角的倒计时——5 分钟。"
        "可以看到我在推挤 AI 的方块，同时 AI 也在反击。"
        "消行时除了之前的黄色闪光，现在还有彩色粒子和屏幕震动，"
        "方块落地的瞬间会有一个白色锁定闪光，让你清楚看到每一块钉死的位置。"
    )
    add_body(doc,
        "另外进入菜单的时候有一个淡入过渡动画，虽然不是游戏核心功能，"
        "但让整个程序的完成度看起来更高了。这些视觉打磨总共不到两百行代码，"
        "但对展示效果的提升是巨大的。"
    )
    add_note(doc, "【录屏放完，切回桌面】")

    add_divider(doc)

    # ==============================================
    add_section(doc, "六、技术实现速览")
    add_note(doc, "【快速带过，不需要展开，给老师一个交代】")

    add_body(doc,
        "技术栈快速说一下：Python 3.10 + Pygame + PyTorch。"
        "DQN 三层全连接 (128,128,64)，输入是 23 维棋盘特征，不是原始图像。"
        "训练 5000 个 episode，PER 加课程学习，模型推理一次不到 1 毫秒。"
        "游戏帧率稳定在 60 FPS，整个项目代码两千多行。"
    )

    add_divider(doc)

    # ==============================================
    add_section(doc, "七、收尾")

    add_body(doc,
        "总结一句话：我这个项目最大的价值不是「实现了一个能玩俄罗斯方块的 AI」，"
        "而是找到了一种方法，让 AI 在保持智能性的同时还能流畅运行，"
        "并且这个方法是可配置、可切换的，不是写死的。"
        "我觉得这是在实时交互系统里落地强化学习算法时真正需要解决的问题。"
    )
    add_body(doc,
        "代码和打包版都会随报告一起提交。欢迎大家课后试玩，看看能不能打赢困难难度的 AI。"
        "谢谢老师和各位！"
    )

    add_divider(doc)

    # ==============================================
    add_section(doc, "附：上台备忘")
    add_body(doc, "1. 带笔记本（教室电脑不一定有 Pygame + PyTorch）")
    add_body(doc, "2. 录屏 MP4 备份在 U 盘里，防止现场跑不起来")
    add_body(doc, "3. 训练曲线截图（图4）和 AI 流程图（图1）各备一张，翻到的时候用")
    add_body(doc, "4. 录屏播放到同屏竞技部分时停下，重点讲碰撞推挤和倒计时")
    add_body(doc, "5. 语速适中，别赶，录屏本身就是最好的说明")

    output_path = "/home/lidio/Tetris_Python/课堂展示发言稿_v2.docx"
    doc.save(output_path)
    print(f"发言稿已保存至: {output_path}")


if __name__ == "__main__":
    main()
